from fastapi import APIRouter, Depends, UploadFile, File, Form, HTTPException, Query, BackgroundTasks
from fastapi.responses import FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload
from typing import List
import os
import re
from html import unescape
import aiofiles
import numpy as np
from docx import Document

from database import get_async_db
from models import Book, Page, OCRResult
from services import handle_pdf_upload, process_document_task
from adapters.factory import AdapterFactory
from utils.math import cosine_similarity
from core.config import settings
from utils.embeddings import get_local_embedding
import asyncio

import datetime
from pydantic import BaseModel

router = APIRouter()

class PagePublishItem(BaseModel):
    page_id: str
    extracted_text: str

class PublishAllPayload(BaseModel):
    pages: List[PagePublishItem]

@router.post("/{book_id}/publish_all")
async def publish_all_pages(
    book_id: str,
    payload: PublishAllPayload,
    db: AsyncSession = Depends(get_async_db)
):
    book = await db.get(Book, book_id)
    if not book:
        raise HTTPException(status_code=404, detail="Book not found")

    # Map the submitted texts by page_id for quick lookup
    text_updates = {item.page_id: item.extracted_text for item in payload.pages}
    
    # Get all pages in this book
    result = await db.execute(select(Page).options(selectinload(Page.ocr_results)).where(Page.book_id == book_id))
    pages = result.scalars().all()
    
    for page in pages:
        if page.id in text_updates:
            new_text = text_updates[page.id]
            
            # Find the latest OCR record, or create it if none exists
            ocr_record = next(iter(page.ocr_results), None)
            
            # Flag to decide if we need to embed
            needs_embedding = False
            
            if not ocr_record:
                # No existing record, create one and embed
                ocr_record = OCRResult(page_id=page.id, extracted_text=new_text)
                db.add(ocr_record)
                needs_embedding = True
            elif ocr_record.extracted_text != new_text:
                # Text changed, update and embed
                ocr_record.extracted_text = new_text
                ocr_record.created_at = datetime.datetime.utcnow()
                needs_embedding = True
                
            if needs_embedding:
                try:
                    embedding_vector = await asyncio.to_thread(get_local_embedding, new_text)
                    if embedding_vector:
                        ocr_record.embedding = np.array(embedding_vector, dtype='float32').tobytes()
                except Exception as e:
                    print(f"Error updating embedding in publish_all: {e}")
                    
        # Mark every page in the payload as Published
        page.status = "Published"
        
    # Mark book as completed/published
    book.status = "Completed"
    
    await db.commit()
    return {"message": "All pages published successfully"}

def _to_plain_text(raw_text: str) -> str:
    # Handle both plain text and simple HTML OCR output.
    text = raw_text or ""
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"</(p|div|h1|h2|h3|h4|h5|h6|li)>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    text = unescape(text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

@router.get("/global/search")
async def global_search(
    query: str, 
    mode: str = Query("semantic", description="Search mode: semantic or keyword"),
    db: AsyncSession = Depends(get_async_db)
):
    if not query:
        return {"results": []}

    # Keyword Search Mode
    if mode == "keyword":
        query_pattern = f"%{query}%"
        result = await db.execute(
            select(Page, OCRResult, Book)
            .join(OCRResult, Page.id == OCRResult.page_id)
            .join(Book, Page.book_id == Book.id)
            .where(OCRResult.extracted_text.like(query_pattern))
        )
        rows = result.all()
        
        results = []
        for page, ocr, book in rows:
            text = _to_plain_text(ocr.extracted_text or "")
            idx = text.lower().find(query.lower())
            start = max(0, idx - 100)
            end = min(len(text), idx + 200)
            snippet = ("..." if start > 0 else "") + text[start:end] + ("..." if end < len(text) else "")
            
            results.append({
                "id": page.id,
                "book_id": book.id,
                "book_title": book.title,
                "page_number": page.page_number,
                "extracted_text": snippet,
                "score": 1.0, 
                "status": page.status
            })
        return {"results": results[:30]}

    # Semantic Search Mode
    try:
        query_embedding = await asyncio.to_thread(get_local_embedding, query)
        if not query_embedding:
            ai_adapter = AdapterFactory.get_adapter(provider=settings.AI_PROVIDER)
            query_embedding = await ai_adapter.get_embedding(query)
            
        if not query_embedding:
            return {"results": [], "error": "Could not generate query embedding"}
        query_vec = np.array(query_embedding, dtype='float32')
    except Exception as e:
        return {"results": [], "error": str(e)}

    result = await db.execute(
        select(Page, OCRResult, Book)
        .join(OCRResult, Page.id == OCRResult.page_id)
        .join(Book, Page.book_id == Book.id)
    )
    rows = result.all()
    
    scored_map = {}
    for page, ocr, book in rows:
        if ocr.embedding:
            page_vec = np.frombuffer(ocr.embedding, dtype='float32')
            score = float(cosine_similarity(query_vec, page_vec))
            
            if score > 0.35:
                if page.id not in scored_map or score > scored_map[page.id]["score"]:
                    text = _to_plain_text(ocr.extracted_text or "")
                    snippet = (text[:300] + "...") if len(text) > 300 else text
                    
                    scored_map[page.id] = {
                        "id": page.id,
                        "book_id": book.id,
                        "book_title": book.title,
                        "page_number": page.page_number,
                        "extracted_text": snippet,
                        "score": score,
                        "status": page.status
                    }
    
    scored_results = list(scored_map.values())
    scored_results.sort(key=lambda x: x["score"], reverse=True)
    return {"results": scored_results[:30]}

@router.get("/")
async def get_all_books(db: AsyncSession = Depends(get_async_db)):
    result = await db.execute(
        select(Book)
        .options(selectinload(Book.pages))
        .order_by(Book.created_at.desc())
    )
    books = result.scalars().all()
    
    response_books = []
    for b in books:
        # Get thumbnail from first page
        thumbnail = None
        if b.pages:
            # Sort pages to ensure we get page 1
            sorted_pages = sorted(b.pages, key=lambda p: p.page_number)
            first_page = sorted_pages[0]
            if first_page.image_path:
                # Convert path to a URL relative to /uploads
                # Assuming image_path stores something like 'uploads/book_id/file.png'
                # or an absolute path. We want the part after 'uploads/'
                path_parts = first_page.image_path.replace("\\", "/").split("/uploads/")
                if len(path_parts) > 1:
                    thumbnail = f"/uploads/{path_parts[1]}"
                else:
                    # Fallback if uploads isn't in the path string (might be a relative path from the start)
                    thumbnail = f"/{first_page.image_path.replace('\\', '/')}"
        
        response_books.append({
            "id": b.id,
            "title": b.title,
            "status": b.status,
            "created_at": b.created_at,
            "thumbnail": thumbnail
        })
        
    return {"books": response_books}

@router.post("/upload")
async def upload_book(
    background_tasks: BackgroundTasks,
    title: str = Form(...),
    files: List[UploadFile] = File(...),
    db: AsyncSession = Depends(get_async_db)
):
    if not files:
        raise HTTPException(status_code=400, detail="No files provided")
        
    book = Book(title=title)
    db.add(book)
    await db.commit()
    await db.refresh(book)
    
    book_folder = os.path.join(settings.UPLOAD_DIR, str(book.id))
    os.makedirs(book_folder, exist_ok=True)
    
    saved_paths = []
    
    # Check if single PDF
    if len(files) == 1 and (files[0].filename.endswith(".pdf") if files[0].filename else False):
        pdf_path = os.path.join(book_folder, (files[0].filename or "document.pdf"))
        async with aiofiles.open(pdf_path, 'wb') as out_file:
            content = await files[0].read()
            await out_file.write(content)
        try:
            saved_paths = await handle_pdf_upload(pdf_path, settings.UPLOAD_DIR)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"PDF conversion failed: {str(e)}")
    else:
        # Multiple images
        for idx, file in enumerate(files):
            file_path = os.path.join(book_folder, (file.filename or f"image_{idx}.png"))
            async with aiofiles.open(file_path, 'wb') as out_file:
                content = await file.read()
                await out_file.write(content)
            saved_paths.append(file_path)
            
    # Create Page records
    for i, path in enumerate(saved_paths):
        page = Page(book_id=book.id, page_number=i+1, image_path=path)
        db.add(page)
        
    await db.commit()
    
    # Automatically start processing in the background after upload
    background_tasks.add_task(process_document_task, f"{book.id}")
    
    return {"message": "Files received successfully and processing started", "book_id": f"{book.id}", "pages_count": len(saved_paths)}

@router.post("/{book_id}/process")
async def process_book(
    book_id: str,
    background_tasks: BackgroundTasks,
    prompt_mode: str | None = Query(None, description="OCR prompt mode override: normal or formatted"),
    db: AsyncSession = Depends(get_async_db),
):
    book = await db.get(Book, book_id)
    if not book:
        raise HTTPException(status_code=404, detail="Book not found")

    background_tasks.add_task(process_document_task, book_id, prompt_mode=prompt_mode)
    return {"message": "Processing started in background", "book_id": book_id, "prompt_mode": prompt_mode or settings.OCR_PROMPT_MODE}

@router.get("/{book_id}/status")
async def get_book_status(book_id: str, db: AsyncSession = Depends(get_async_db)):
    book = await db.get(Book, book_id)
    if not book:
        raise HTTPException(status_code=404, detail="Book not found")
        
    result = await db.execute(select(Page).where(Page.book_id == book_id))
    pages = result.scalars().all()
    
    completed = sum(1 for p in pages if p.status in ["Completed", "Published"])
    total = len(pages)
    progress_percent = (completed / total * 100) if total > 0 else 0
    
    pages_status = [{"page_number": p.page_number, "status": p.status, "id": p.id} for p in pages]
    
    return {
        "status": book.status,
        "title": book.title,
        "progress": f"{completed}/{total}",
        "progress_percent": progress_percent,
        "pages": pages_status
    }

@router.get("/{book_id}/results")
async def get_book_results(book_id: str, db: AsyncSession = Depends(get_async_db)):
    result = await db.execute(
        select(Page)
        .where(Page.book_id == book_id)
        .options(selectinload(Page.ocr_results))
        .order_by(Page.page_number)
    )
    pages = result.scalars().all()
    
    results = []
    for page in pages:
        ocr_data = page.ocr_results[0] if page.ocr_results else None
        text = ocr_data.extracted_text if ocr_data else ""
        confidence = ocr_data.confidence_score if ocr_data else 0.0
        
        results.append({
            "id": page.id,
            "page_number": page.page_number,
            "image_path": page.image_path,
            "extracted_text": text,
            "confidence": confidence,
            "status": page.status
        })
        
    book = await db.get(Book, book_id)
    return {
        "book_id": book_id, 
        "title": book.title if book else "كتاب غير معروف",
        "results": results
    }

@router.get("/{book_id}/search")
async def search_book(
    book_id: str, 
    query: str, 
    mode: str = Query("semantic", description="Search mode: semantic or keyword"),
    db: AsyncSession = Depends(get_async_db)
):
    if not query:
        return {"results": []}

    # Keyword Search Mode (Exact/Partial phrase matching)
    if mode == "keyword":
        # Search using SQL LIKE for simple keyword matching
        query_pattern = f"%{query}%"
        result = await db.execute(
            select(Page, OCRResult)
            .join(OCRResult, Page.id == OCRResult.page_id)
            .where(Page.book_id == book_id)
            .where(OCRResult.extracted_text.like(query_pattern))
            .order_by(Page.page_number)
        )
        rows = result.all()
        
        results = []
        for page, ocr in rows:
            # Basic snippet generation
            text = _to_plain_text(ocr.extracted_text or "")
            idx = text.lower().find(query.lower())
            # Find a good window around the match
            start = max(0, idx - 100)
            end = min(len(text), idx + 200)
            snippet = ("..." if start > 0 else "") + text[start:end] + ("..." if end < len(text) else "")
            
            results.append({
                "id": page.id,
                "page_number": page.page_number,
                "extracted_text": snippet,
                "score": 1.0, # Exact/Keyword matches get a flat score
                "status": page.status
            })
        return {"results": results[:20]}

    # Semantic Search Mode (Vector Similarity)
    try:
        # Use local embedding for the query (same as used for indexing)
        query_embedding = await asyncio.to_thread(get_local_embedding, query)
        if not query_embedding:
            # Fallback (optional)
            ai_adapter = AdapterFactory.get_adapter(provider=settings.AI_PROVIDER)
            query_embedding = await ai_adapter.get_embedding(query)
            
        if not query_embedding:
            return {"results": [], "error": "Could not generate query embedding (local or remote)"}
        query_vec = np.array(query_embedding, dtype='float32')
    except Exception as e:
        return {"results": [], "error": str(e)}

    result = await db.execute(
        select(Page, OCRResult)
        .join(OCRResult, Page.id == OCRResult.page_id)
        .where(Page.book_id == book_id)
    )
    rows = result.all()
    
    scored_map = {}
    for page, ocr in rows:
        if ocr.embedding:
            page_vec = np.frombuffer(ocr.embedding, dtype='float32')
            score = float(cosine_similarity(query_vec, page_vec))
            
            # Use a slightly lower threshold for semantic discovery
            if score > 0.35:
                if page.id not in scored_map or score > scored_map[page.id]["score"]:
                    # Basic snippet for semantic - just start of plain text
                    text = _to_plain_text(ocr.extracted_text or "")
                    snippet = (text[:300] + "...") if len(text) > 300 else text
                    
                    scored_map[page.id] = {
                        "id": page.id,
                        "page_number": page.page_number,
                        "extracted_text": snippet,
                        "score": score,
                        "status": page.status
                    }
    
    scored_results = list(scored_map.values())
    scored_results.sort(key=lambda x: x["score"], reverse=True)
    return {"results": scored_results[:15]}

@router.get("/{book_id}/export")
async def export_book_results(book_id: str, db: AsyncSession = Depends(get_async_db)):
    result = await db.execute(
        select(Page)
        .where(Page.book_id == book_id)
        .options(selectinload(Page.ocr_results))
        .order_by(Page.page_number)
    )
    pages = result.scalars().all()

    book = await db.get(Book, book_id)
    book_title = book.title if book and book.title else f"book_{book_id}"

    document = Document()
    document.add_heading(book_title, level=1)

    for page in pages:
        ocr_data = page.ocr_results[0] if page.ocr_results else None
        if ocr_data and ocr_data.extracted_text:
            document.add_heading(f"صفحة {page.page_number}", level=2)
            plain_text = _to_plain_text(ocr_data.extracted_text)
            if plain_text:
                for paragraph in plain_text.split("\n"):
                    cleaned = paragraph.strip()
                    if cleaned:
                        document.add_paragraph(cleaned)
            else:
                document.add_paragraph("")

            document.add_page_break()

    export_path = os.path.join(settings.UPLOAD_DIR, f"{book_id}_export.docx")
    document.save(export_path)

    safe_name = re.sub(r"[^\w\u0600-\u06FF\s-]", "", book_title).strip() or f"exported_book_{book_id}"
    return FileResponse(
        export_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{safe_name}.docx",
    )
